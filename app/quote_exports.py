"""DOCX and PDF generated from exactly the same persisted calculation."""

import io
from pathlib import Path
from xml.sax.saxutils import escape

from app.errors import DomainError

FONT_PATH = Path(__file__).parent / "assets" / "fonts" / "NotoSans-Regular.ttf"


def content(full):
    revision = full["revision"]
    snapshot = revision["snapshot"]
    draft = revision["status"] != "approved"
    title = ("ЧЕРНОВИК / DRAFT — " if draft else "") + snapshot["input"]["title"]
    paragraphs = [
        f"AI Office · КП / Quote {full['id']} · версия / version {revision['version']}",
        f"Заказчик / Customer: {snapshot['input']['customer']}",
        snapshot["input"]["accompanying_text"],
        f"Прайс / Catalog: {snapshot['catalog_name']} · версия {snapshot['catalog_version']}",
    ]
    source = snapshot["source"]
    if source:
        paragraphs.append(f"Заявка / Source: {source['name']} · версия {source['version']}")
    calculation = snapshot["calculation"]
    table = [
        [
            "Артикул / SKU",
            "Наименование / Item",
            "Кол-во / Qty",
            "Цена без НДС / Price",
            "Скидка %",
            "НДС / VAT",
            "Итого / Total",
        ]
    ]
    for item in calculation["lines"]:
        table.append(
            [
                item["sku"],
                item["name"],
                f"{item['quantity']} {item['unit']}",
                item["price_without_vat"],
                item["discount_percent"],
                f"{item['vat']} ({item['vat_percent']}%)",
                item["total"],
            ]
        )
    totals = f"Без НДС / Net: {calculation['net']} RUB · НДС / VAT: {calculation['vat']} RUB · Итого / Total: {calculation['total']} RUB"
    provenance = [
        f"Расчёт / Calculation: {revision['content_hash']}",
        "Округление каждой строки до копеек / Round each line to 0.01: ROUND_HALF_UP.",
    ]
    for item in calculation["lines"]:
        evidence = item["evidence"]
        provenance.append(
            f"{item['sku']}: net = {item['price_without_vat']} × {item['quantity']} × (1 - {item['discount_percent']} / 100) = {item['net']}; VAT = {item['net']} × {item['vat_percent']} / 100 = {item['vat']}. Прайс / catalog {evidence['catalog_version_id']}, строка / row {evidence['source_row']}, SHA-256 {evidence['content_hash']}."
        )
    return title, paragraphs, table, totals, provenance


def export_quote(full, format_name):
    if format_name not in {"docx", "pdf"}:
        raise DomainError("UNSUPPORTED_QUOTE_EXPORT", 415)
    title, paragraphs, table, totals, provenance = content(full)
    output = io.BytesIO()
    if format_name == "docx":
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Mm(297), Mm(210)
        section.left_margin = section.right_margin = Pt(22)
        section.top_margin = section.bottom_margin = Pt(25)
        for style_name in ("Normal", "Title", "Heading 1"):
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.underline = False
            for border in style.element.xpath(".//w:pBdr"):
                border.getparent().remove(border)
        document.styles["Normal"].font.size = Pt(10)
        document.styles["Normal"].paragraph_format.space_after = Pt(7)
        document.styles["Title"].font.size = Pt(20)
        document.styles["Title"].paragraph_format.space_after = Pt(12)
        document.styles["Heading 1"].font.size = Pt(12)
        document.add_heading(title, 0)
        for text in paragraphs:
            document.add_paragraph(text)
        grid = document.add_table(rows=0, cols=len(table[0]))
        grid.style = "Table Grid"
        grid.autofit = False
        widths = [82, 203, 65, 83, 60, 100, 105]
        for column, width in zip(grid.columns, widths, strict=True):
            column.width = Pt(width)
        for row_index, values in enumerate(table):
            table_row = grid.add_row()
            table_row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            if row_index == 0:
                table_row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
            for cell, value, width in zip(table_row.cells, values, widths, strict=True):
                cell.width = Pt(width)
                cell.text = value
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(4)
                    paragraph.paragraph_format.space_before = Pt(4)
                    for run in paragraph.runs:
                        run.font.size = Pt(8 if width == 105 and len(value) > 15 else 9)
                        run.bold = row_index == 0
                if row_index == 0:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "E8EDF3")
                    cell._tc.get_or_add_tcPr().append(shading)
        document.add_paragraph(totals)
        document.add_heading("Основание расчёта / Calculation evidence", 1)
        for text in provenance:
            paragraph = document.add_paragraph(text)
            for run in paragraph.runs:
                run.font.size = Pt(8)
        document.save(output)
    else:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import LongTable, Paragraph, SimpleDocTemplate, Spacer, TableStyle

        if "AI-Office-Noto" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("AI-Office-Noto", str(FONT_PATH)))
        normal = ParagraphStyle(
            "quote", fontName="AI-Office-Noto", fontSize=9, leading=13, spaceAfter=7, wordWrap="CJK"
        )
        heading = ParagraphStyle("heading", parent=normal, fontSize=17, leading=22, spaceAfter=14)
        small = ParagraphStyle("small", parent=normal, fontSize=7, leading=10)
        para = lambda value, style=normal: Paragraph(escape(value).replace("\n", "<br/>"), style)  # noqa: E731
        flow = [para(title, heading), *(para(value) for value in paragraphs), Spacer(1, 7)]
        grid = LongTable(
            [[para(value, small) for value in values] for values in table],
            colWidths=[82, 203, 65, 83, 60, 100, 105],
            repeatRows=1,
            hAlign="LEFT",
        )
        grid.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#b6c1cc")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]
            )
        )
        flow += [grid, Spacer(1, 12), para(totals), *(para(value, small) for value in provenance)]
        document = SimpleDocTemplate(
            output,
            pagesize=landscape(A4),
            leftMargin=22,
            rightMargin=22,
            topMargin=25,
            bottomMargin=25,
            title=title,
            author="AI Office",
        )
        document.build(flow)
    return output.getvalue()
