"""Exercise the customer journey in a real browser using explicitly synthetic sample files."""

import argparse
import json
import time
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from playwright.sync_api import expect, sync_playwright
from pypdf import PdfReader
from reportlab.pdfgen import canvas
from smoke import BASE, demo_token

EXAMPLES = Path("examples/customer-demo")


def prepare_examples():
    EXAMPLES.mkdir(parents=True, exist_ok=True)
    (EXAMPLES / "synthetic-catalog.csv").write_bytes(
        Path("examples/catalogs/synthetic-demo.csv").read_bytes()
    )
    document = Document()
    document.add_heading("SYNTHETIC CUSTOMER REQUEST / УЧЕБНАЯ ЗАЯВКА", 0)
    document.add_paragraph("STEEL-01 × 3\nBOLT-01 × 10")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Условие"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Срок поставки"
    table.cell(1, 1).text = "Три рабочих дня — синтетический пример"
    document.save(EXAMPLES / "synthetic-request.docx")
    pdf = canvas.Canvas(str(EXAMPLES / "synthetic-terms.pdf"), invariant=True)
    pdf.drawString(50, 780, "SYNTHETIC CUSTOMER TERMS - NOT A REAL CONTRACT")
    pdf.drawString(50, 750, "Orchid customer delivery clause: three working days.")
    pdf.showPage()
    pdf.drawString(50, 780, "SYNTHETIC SECOND PAGE")
    pdf.drawString(50, 750, "Steel goods require a written quality certificate.")
    pdf.save()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--prepare-only", action="store_true")
    args = parser.parse_args()
    if args.prepare_only or not all(
        (EXAMPLES / name).exists()
        for name in ("synthetic-catalog.csv", "synthetic-request.docx", "synthetic-terms.pdf")
    ):
        prepare_examples()
    if args.prepare_only:
        print("Prepared explicitly synthetic CSV, DOCX and PDF examples.")
        return
    output = Path(".runtime/customer-demo")
    output.mkdir(parents=True, exist_ok=True)
    screenshots = Path("docs/assets")
    screenshots.mkdir(parents=True, exist_ok=True)
    errors, external, checks = [], [], []
    with sync_playwright() as p:
        chrome = Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome")
        browser = p.chromium.launch(
            headless=True, executable_path=str(chrome) if chrome.exists() else None
        )
        context = browser.new_context(
            viewport={"width": 1440, "height": 1050},
            timezone_id="Europe/Moscow",
            accept_downloads=True,
        )

        def route(request_route):
            if not request_route.request.url.startswith(BASE + "/"):
                external.append(request_route.request.url)
                request_route.abort()
            else:
                # Deliberately delay history reads: old implementations cleared selects between
                # awaited requests and changed approval inputs while the user was editing.
                if (
                    "/api/v1/catalogs/" in request_route.request.url
                    and "/versions" in request_route.request.url
                ):
                    time.sleep(0.03)
                request_route.continue_()

        context.route("**/*", route)
        page = context.new_page()
        page.on("pageerror", lambda error: errors.append(str(error)))
        page.goto(BASE)
        page.locator("#token").fill(demo_token())
        page.locator("#login-form button").click()
        page.locator("#workspace").wait_for(state="visible")
        page.locator("#engine-label").filter(has_text="deterministic_demo").wait_for()
        page.locator('nav a[data-page="knowledge"]').click()
        for name in ("synthetic-request.docx", "synthetic-terms.pdf"):
            page.locator("#file").set_input_files(EXAMPLES / name)
            with page.expect_response(
                lambda response: (
                    response.url.endswith("/api/v1/documents") and response.request.method == "POST"
                )
            ) as uploaded_document:
                page.locator("#upload-form button").click()
            assert uploaded_document.value.status == 200
            page.locator("#document-list .row").filter(has_text=name).wait_for()
        page.locator("#question").fill("Orchid customer delivery clause")
        page.locator("#ask-form button").click()
        page.locator("#answer-result .evidence").filter(has_text="Orchid customer").first.wait_for(
            timeout=45000
        )
        page.locator("#answer-result .evidence").filter(has_text="Orchid customer").first.locator(
            "button"
        ).click()
        page.locator(".source-focus").wait_for()
        assert "three working days" in page.locator(".source-focus").inner_text()
        page.screenshot(path=str(screenshots / "customer-documents.png"), full_page=True)
        page.locator("#close-detail").click()
        checks.append("PDF/DOCX upload, page citation and highlighted source")
        page.locator('nav a[data-page="quotes"]').click()
        page.locator("#catalog-file").set_input_files(EXAMPLES / "synthetic-catalog.csv")
        with page.expect_response(
            lambda response: (
                response.url.endswith("/api/v1/catalogs") and response.request.method == "POST"
            )
        ) as uploaded_catalog:
            page.locator("#catalog-form button[type=submit]").click()
        assert uploaded_catalog.value.status == 200
        catalog_version_id = uploaded_catalog.value.json()["id"]
        page.locator("#quote-catalog option").filter(
            has_text="synthetic-catalog.csv"
        ).first.wait_for(state="attached")
        # Choose the just-uploaded current version; retained catalogs from earlier runs are harmless.
        page.wait_for_function(
            "id => document.querySelector('#quote-catalog').value === id", arg=catalog_version_id
        )
        source_value = (
            page.locator("#quote-source option")
            .filter(has_text="synthetic-request.docx")
            .first.get_attribute("value")
        )
        page.locator("#quote-source").select_option(source_value)
        with page.expect_response(
            lambda response: response.url.endswith("/api/v1/quote-suggestions")
        ) as prepared:
            page.locator("#quote-request-form button[type=submit]").click()
        assert prepared.value.status == 202, prepared.value.text()
        try:
            page.locator("#quote-suggestion .status.completed").wait_for(timeout=45000)
        except Exception:
            print(
                json.dumps(
                    {
                        "notice": page.locator("#notice").inner_text(),
                        "suggestion_state": page.locator("#quote-suggestion").inner_text(),
                        "browser_errors": errors,
                    }
                )
            )
            raise
        expect(page.locator(".quote-sku")).to_have_count(2)
        title = "Synthetic customer walkthrough " + str(int(time.time()))
        page.locator("#quote-title").fill(title)
        page.locator("#quote-customer").fill("Учебная компания / Synthetic customer")
        page.locator(".quote-discount").first.fill("5")
        page.locator("#quote-assignee").select_option("employee")
        page.locator("#quote-due").fill(
            (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%dT12:00")
        )
        page.locator("#save-quote").click()
        page.locator("#quote-result .quote-total").wait_for()
        expect(page.locator("#quote-history")).to_have_value("1")
        with page.expect_download() as draft_info:
            page.locator('#quote-result a[download="quote.pdf"]').click()
        draft = output / "draft.pdf"
        draft_info.value.save_as(draft)
        assert "DRAFT" in " ".join(p.extract_text() for p in PdfReader(draft).pages)
        page.locator(".quote-quantity").first.fill("4")
        page.locator("#save-quote").click()
        expect(page.locator("#quote-history")).to_have_value("2")
        page.locator("#quote-history").select_option("1")
        expect(page.locator("#save-quote")).to_be_disabled()
        page.locator("#quote-history").select_option("2")
        expect(page.locator("#save-quote")).to_be_enabled()
        page.locator("#quote-result button").filter(has_text="Why?").click()
        page.locator("#detail").wait_for(state="visible")
        assert "100.00" in page.locator("#detail-content").inner_text()
        page.locator("#close-detail").click()
        page.locator("#quote-result button").filter(has_text="Request approval").click()
        page.locator('#quote-result button[data-decision="approve"]').wait_for()
        # Unsaved changes cannot approve the previous calculation; saving supersedes it.
        page.locator("#quote-text").fill("Reviewed synthetic proposal; customer feedback required.")
        expect(page.locator('#quote-result button[data-decision="approve"]')).to_be_disabled()
        page.locator("#save-quote").click()
        expect(page.locator("#quote-history")).to_have_value("3")
        page.locator("#quote-result button").filter(has_text="Request approval").click()
        page.locator('#quote-result button[data-decision="approve"]').wait_for()
        # Simulate overlapping background refreshes while the catalog history responses lag.
        selected = page.locator("#quote-catalog").input_value()
        page.evaluate("Promise.all([loadQuotes(), loadQuotes()])")
        expect(page.locator("#quote-catalog")).to_have_value(selected)
        expect(page.locator('#quote-result button[data-decision="approve"]')).to_be_enabled()
        page.screenshot(path=str(screenshots / "customer-quotes.png"), full_page=True)
        page.locator('#quote-result button[data-decision="approve"]').click()
        page.locator("#quote-result .success").wait_for(timeout=45000)
        for format in ("pdf", "docx"):
            with page.expect_download() as export_info:
                page.locator(f'#quote-result a[download="quote.{format}"]').click()
            export_info.value.save_as(output / f"approved.{format}")
        assert "DRAFT" not in " ".join(
            p.extract_text() for p in PdfReader(output / "approved.pdf").pages
        )
        assert "606.00" in " ".join(
            p.extract_text() for p in PdfReader(output / "approved.pdf").pages
        )
        checks.append(
            "CSV catalog, document suggestion, edited quote, immutable revisions, formula and draft/approved exports"
        )
        page.locator("#quote-result button").filter(has_text="Open follow-up task").click()
        expect(page.locator(".task-card")).to_have_count(1)
        assert "employee" in page.locator(".task-card").inner_text()
        page.locator(".task-card select").select_option("blocked")
        page.locator(".task-card textarea").fill(
            "Synthetic follow-up: waiting for customer confirmation."
        )
        page.locator(".task-card button[type=submit]").click()
        page.locator('#task-filters [data-filter="overdue"]').click()
        page.locator(".task-card .status.overdue").first.wait_for()
        page.locator('#task-filters [data-filter="blocked"]').click()
        page.locator(".task-card .status.blocked").first.wait_for()
        page.screenshot(path=str(screenshots / "customer-tasks.png"), full_page=True)
        page.locator('nav a[data-page="today"]').click()
        page.locator("#refresh-briefing").click()
        page.locator("#briefing-result").filter(has_text="Overdue").wait_for()
        checks.append(
            "Approved personal assignment, exact task link, overdue/blocked filters and current briefing"
        )
        page.locator('nav a[data-page="quotes"]').click()
        page.locator("#lang").click()
        expect(page.locator("#page-quotes h1")).to_have_text("Коммерческие предложения")
        expect(page.locator("#quote-suggestion .status.completed")).to_have_text("Завершено")
        expect(page.locator("#quote-result .status.approved")).to_have_text("Согласовано")
        expect(page.locator("#quote-lines button").first).to_have_text("Удалить")
        assert "Часовой пояс" in page.locator("#quote-editor .timezone-note").inner_text()
        page.screenshot(path=str(screenshots / "customer-quotes-ru.png"), full_page=True)
        page.set_viewport_size({"width": 390, "height": 844})
        page.screenshot(path=str(screenshots / "customer-mobile.png"), full_page=True)
        assert page.evaluate("document.documentElement.scrollWidth <= innerWidth")
        checks.append("Russian/English and mobile layout")
        browser.close()
    assert not errors, errors
    assert not external, external
    report = {
        "checks": checks,
        "browser_errors": errors,
        "external_requests": external,
        "screenshots": 5,
        "data": "explicitly synthetic customer examples",
    }
    Path(".runtime/customer-demo.json").write_text(json.dumps(report, ensure_ascii=False, indent=2))
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
