import io
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfReader
from sqlalchemy import select

from app.contracts import Decision
from app.db import row
from app.errors import DomainError
from app.providers import DemoProvider
from app.quote_contracts import QuoteDraft, QuoteLine
from app.quotes import Quotes, calculate, parse_catalog
from app.schema_v1 import tasks

CSV = Path("examples/catalogs/synthetic-demo.csv").read_bytes()


def service(ctx):
    return Quotes(ctx["engine"], ctx["settings"], ctx["knowledge"], ctx["clock"])


def draft(catalog, **changes):
    return QuoteDraft.model_validate(
        {
            "catalog_version_id": catalog["id"],
            "title": "КП на сталь",
            "customer": "Компания Пример",
            "lines": [{"sku": "STEEL-01", "quantity": "3", "discount_percent": "5"}],
            "task": {
                "title": "Проверить КП",
                "team_id": "procurement",
                "assignee_id": "employee",
                "due_at": "2026-09-11T15:00:00+03:00",
                "acceptance_criteria": "Проверены состав и итог",
            },
            **changes,
        }
    )


def prepared(ctx):
    quotes = service(ctx)
    actor = ctx["actors"]["owner"]
    catalog = quotes.import_catalog(actor, "catalog.csv", CSV, "test")
    return quotes, actor, catalog, quotes.save(actor, draft(catalog), "test")


def test_decimal_rounding_and_trace():
    catalog = {
        "id": "catalog-v1",
        "version": 1,
        "content_hash": "hash",
        "rows": [
            {
                "sku": "X",
                "name": "Test",
                "unit": "шт",
                "price_without_vat": "0.05",
                "vat_percent": "10",
                "source_row": 2,
            }
        ],
    }
    result = calculate([QuoteLine(sku="X", quantity=Decimal(1))], catalog)
    assert result["net"] == "0.05"
    assert result["vat"] == "0.01"  # Half-up, not binary/banker's rounding.
    assert result["total"] == "0.06"
    assert result["lines"][0]["evidence"]["source_row"] == 2


def test_catalog_versions_do_not_change_old_quote(ctx):
    quotes, actor, catalog, quote = prepared(ctx)
    second = quotes.import_catalog(
        actor, "catalog.csv", CSV.replace(b"100.00", b"200.00"), "test", catalog["catalog_id"]
    )
    assert second["version"] == 2
    assert quotes.catalog(actor, catalog["id"])["current"] is False
    assert (
        quotes.get(actor, quote["id"])["revision"]["snapshot"]["calculation"]["total"] == "342.00"
    )
    assert quotes.import_catalog(
        actor, "catalog.csv", CSV.replace(b"100.00", b"200.00"), "test", catalog["catalog_id"]
    )["replayed"]


@pytest.mark.parametrize(
    "replace,code",
    [
        (b"=2+2", "CATALOG_FORMULA_NOT_ALLOWED"),
        (b"NaN", "INVALID_CATALOG_CELL"),
        (b"-1", "INVALID_CATALOG_CELL"),
        (b"10.001", "INVALID_CATALOG_CELL"),
    ],
)
def test_bad_csv_price_rejected(replace, code):
    with pytest.raises(DomainError) as caught:
        parse_catalog("catalog.csv", CSV.replace(b"100.00", replace))
    assert caught.value.code == code


def test_duplicate_and_currency_rejected():
    with pytest.raises(DomainError, match="DUPLICATE_SKU"):
        parse_catalog("catalog.csv", CSV.replace(b"BOLT-01", b"STEEL-01"))
    with pytest.raises(DomainError, match="CATALOG_CURRENCY_RUB_REQUIRED"):
        parse_catalog("catalog.csv", CSV.replace(b"RUB", b"USD"))


def xlsx(formula=False):
    wb = Workbook()
    sheet = wb.active
    sheet.append(["sku", "name", "unit", "price_without_vat", "vat_percent", "currency"])
    sheet.append(["X", "Пример", "шт", "=2+2" if formula else 2.5, 20, "RUB"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_xlsx_values_and_formula_cell():
    assert parse_catalog("prices.xlsx", xlsx())[0]["price_without_vat"] == "2.50"
    with pytest.raises(DomainError) as caught:
        parse_catalog("prices.xlsx", xlsx(True))
    assert caught.value.code == "CATALOG_FORMULA_NOT_ALLOWED"
    assert "D2" in caught.value.message


def test_unknown_sku_and_discount(ctx):
    quotes, actor, catalog, _ = prepared(ctx)
    with pytest.raises(DomainError, match="UNKNOWN_SKU"):
        quotes.save(actor, draft(catalog, lines=[{"sku": "UNKNOWN", "quantity": "1"}]), "test")
    result = quotes.save(
        actor,
        draft(catalog, lines=[{"sku": "STEEL-01", "quantity": "3", "discount_percent": "100"}]),
        "test",
    )
    assert result["revision"]["snapshot"]["calculation"]["total"] == "0.00"


def test_export_same_snapshot_cyrillic_and_draft(ctx):
    quotes, actor, _, quote = prepared(ctx)
    pdf = quotes.export(actor, quote["id"], "pdf")
    docx = quotes.export(actor, quote["id"], "docx")
    text = "\n".join(page.extract_text() for page in PdfReader(io.BytesIO(pdf)).pages)
    word = Document(io.BytesIO(docx))
    word_text = "\n".join(p.text for p in word.paragraphs)
    assert "ЧЕРНОВИК" in text and "ЧЕРНОВИК" in word_text
    assert "Компания Пример" in text and "Компания Пример" in word_text
    assert "342.00" in text and "342.00" in word_text
    assert "1 - 5 / 100" in text
    assert quote["revision"]["content_hash"] in text
    assert "/FontFile2" in str(
        PdfReader(io.BytesIO(pdf)).pages[0]["/Resources"]["/Font"]["/F2+0"]["/FontDescriptor"]
    )


def test_revision_invalidates_previous_approval(ctx):
    quotes, actor, catalog, quote = prepared(ctx)
    proposed = quotes.propose(actor, quote["id"], 1, "quote-propose", "test")
    proposal = proposed["proposal"]
    decision = Decision(decision="approve", version=1, payload_hash=proposal["payload_hash"])
    ctx["work"].decide(actor, proposal["id"], decision)
    updated = quotes.save(
        actor, draft(catalog, customer="Изменённый заказчик"), "test", quote["id"], 1
    )
    assert updated["revision"]["version"] == 2
    assert ctx["work"].process_one() is False
    with pytest.raises(DomainError):
        ctx["work"].decide(actor, proposal["id"], decision)
    with ctx["engine"].connect() as conn:
        assert row(conn, select(tasks)) is None


def test_approved_quote_creates_exactly_one_task(ctx):
    quotes, actor, _, quote = prepared(ctx)
    proposed = quotes.propose(actor, quote["id"], 1, "quote-propose", "test")
    proposal = proposed["proposal"]
    decision = Decision(decision="approve", version=1, payload_hash=proposal["payload_hash"])
    ctx["work"].decide(actor, proposal["id"], decision)
    assert ctx["work"].process_one()
    assert ctx["work"].get(actor, proposed["run"]["id"])["state"] == "completed"
    assert ctx["work"].decide(actor, proposal["id"], decision)["replayed"]
    assert not ctx["work"].process_one()
    result = quotes.get(actor, quote["id"])
    assert result["revision"]["status"] == "approved"
    pdftext = "\n".join(
        p.extract_text()
        for p in PdfReader(io.BytesIO(quotes.export(actor, quote["id"], "pdf"))).pages
    )
    assert "ЧЕРНОВИК" not in pdftext
    with ctx["engine"].connect() as conn:
        created = conn.execute(select(tasks)).mappings().all()
    assert len(created) == 1 and created[0]["title"] == "Проверить КП"


def test_demo_suggestions_are_explicit_and_reviewable(ctx):
    quotes, actor, catalog, _ = prepared(ctx)
    result = quotes.suggest(
        actor,
        {"catalog_version_id": catalog["id"], "request_text": "STEEL-01 × 12,5; BOLT-01 x 20"},
        DemoProvider(),
    )
    assert result["lines"][0]["quantity"] == "12.5"
    assert len(result["lines"]) == 2
    assert result["synthetic"] and result["requires_review"] and result["unresolved"]


def test_model_unknown_sku_is_rejected(ctx):
    quotes, actor, catalog, _ = prepared(ctx)

    class Model:
        engine, model_id = "crewai_local", "fixture-model"

        def suggest_quote(self, text, rows):
            return {"lines": [{"sku": "HALLUCINATED", "quantity": "1"}]}

    with pytest.raises(DomainError, match="UNKNOWN_SKU"):
        quotes.suggest(
            actor, {"catalog_version_id": catalog["id"], "request_text": "items"}, Model()
        )


def test_claimed_quote_recovers_after_worker_restart(ctx):
    from app.workflows import Workflows

    quotes, actor, _, quote = prepared(ctx)
    proposed = quotes.propose(actor, quote["id"], 1, "restart-quote", "test")
    proposal = proposed["proposal"]
    ctx["work"].decide(
        actor,
        proposal["id"],
        Decision(decision="approve", version=1, payload_hash=proposal["payload_hash"]),
    )
    abandoned_job = ctx["work"].claim()
    ctx["clock"].value += ctx["settings"].lease_seconds + 1
    replacement = Workflows(
        ctx["engine"], ctx["settings"], ctx["work"].provider, ctx["knowledge"], ctx["clock"]
    )
    assert replacement.process_one()
    assert not replacement.process_one()
    ctx["work"].execute(abandoned_job)  # Fenced stale worker cannot duplicate committed effects.
    with ctx["engine"].connect() as conn:
        created = conn.execute(select(tasks)).mappings().all()
    assert len(created) == 1
    assert quotes.get(actor, quote["id"])["revision"]["status"] == "approved"


def test_concurrent_duplicate_quote_approval(ctx):
    from concurrent.futures import ThreadPoolExecutor

    quotes, actor, _, quote = prepared(ctx)
    proposed = quotes.propose(actor, quote["id"], 1, "concurrent-quote", "test")
    proposal = proposed["proposal"]
    decision = Decision(decision="approve", version=1, payload_hash=proposal["payload_hash"])
    with ThreadPoolExecutor(max_workers=2) as executor:
        results = list(
            executor.map(lambda _: ctx["work"].decide(actor, proposal["id"], decision), range(2))
        )
    assert sorted(result["replayed"] for result in results) == [False, True]
    assert ctx["work"].process_one()
    assert not ctx["work"].process_one()
    with ctx["engine"].connect() as conn:
        assert len(conn.execute(select(tasks)).all()) == 1


def test_untrusted_xlsx_dimensions_cannot_hide_prices():
    import zipfile

    original = xlsx()
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(original)) as source, zipfile.ZipFile(output, "w") as target:
        for name in source.namelist():
            content = source.read(name)
            if name == "xl/worksheets/sheet1.xml":
                content = content.replace(b'ref="A1:F2"', b'ref="A1:F1"')
            target.writestr(name, content)
    assert len(parse_catalog("catalog.xlsx", output.getvalue())) == 1


def test_quote_xml_control_character_rejected():
    from pydantic import ValidationError

    with pytest.raises(ValidationError):
        draft({"id": "catalog-v1"}, customer="Name\x00Hidden")


def test_long_quote_exports_escape_markup_and_preserve_all_lines(ctx):
    quotes = service(ctx)
    actor = ctx["actors"]["owner"]
    csvtext = "sku,name,unit,price_without_vat,vat_percent,currency\n" + "\n".join(
        f"LONG-{i}," + "Длинное наименование <b> & проверка " * 5 + ",шт,999999999.99,20,RUB"
        for i in range(100)
    )
    catalog = quotes.import_catalog(actor, "long.csv", csvtext.encode(), "test")
    quote = quotes.save(
        actor,
        draft(
            catalog,
            customer="Покупатель <b> & заказчик " * 10,
            lines=[
                {"sku": f"LONG-{i}", "quantity": "1000000", "discount_percent": "0"}
                for i in range(100)
            ],
        ),
        "test",
    )
    pdf = quotes.export(actor, quote["id"], "pdf")
    docx = quotes.export(actor, quote["id"], "docx")
    text = "\n".join(page.extract_text() for page in PdfReader(io.BytesIO(pdf)).pages)
    assert "LONG-99" in text and "&" in text
    assert len(Document(io.BytesIO(docx)).tables[0].rows) == 101


@pytest.mark.parametrize("engine", ["deterministic_demo", "crewai_local"])
def test_oversized_request_document_never_returns_partial_suggestions(ctx, engine):
    quotes, actor, catalog, _ = prepared(ctx)
    source = ctx["knowledge"].import_document(
        actor,
        "long-request.txt",
        ("STEEL-01 x 1\n" + "A" * 16500 + "\nBOLT-01 x 9").encode(),
        ["owner"],
        "test",
    )

    class Provider:
        model_id = "fixture"

        def suggest_quote(self, text, rows):
            pytest.fail("Oversized request reached the model")

    provider = Provider()
    provider.engine = engine
    with pytest.raises(DomainError, match="REQUEST_TOO_LARGE_FOR_SUGGESTION"):
        quotes.suggest(
            actor,
            {
                "catalog_version_id": catalog["id"],
                "source_document_id": source["document_id"],
                "source_document_version": source["version"],
            },
            provider,
        )
