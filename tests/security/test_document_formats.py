"""Real binary extraction, immutable originals and permission boundaries."""

import io
import subprocess
import zipfile

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from sqlalchemy import select

from app.contracts import DocumentACL
from app.db import row, transaction
from app.document_parser import parse_document
from app.errors import DomainError
from app.schema_v1 import versions


def pdf_bytes(text="Orchid supply price 125 RUB."):
    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, invariant=True)
    if text:
        pdf.drawString(50, 780, text)
    pdf.showPage()
    pdf.save()
    return stream.getvalue()


def docx_bytes(text="Правила закупок: лимит 120 000 рублей."):
    doc = Document()
    doc.add_paragraph(text)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Артикул"
    table.cell(0, 1).text = "Цена"
    table.cell(1, 0).text = "STEEL-01"
    table.cell(1, 1).text = "125.00"
    doc.add_paragraph("Условия поставки: три дня.")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_pdf_page_anchors_and_word_table_order():
    pdf = parse_document("price.pdf", pdf_bytes(), "application/pdf")
    assert pdf["anchors"] == [
        {"ref": "page:1", "label": "Page 1", "start": 0, "end": len(pdf["text"])}
    ]
    word = parse_document("rules.docx", docx_bytes())
    assert [a["ref"] for a in word["anchors"]] == [
        "paragraph:1",
        "table:1:row:1",
        "table:1:row:2",
        "paragraph:2",
    ]
    assert "120 000 рублей" in word["text"]
    assert "STEEL-01 | 125.00" in word["text"]
    for anchor in word["anchors"]:
        assert word["text"][anchor["start"] : anchor["end"]].strip()


def test_pdf_encryption_scanned_and_corruption_errors():
    writer = PdfWriter()
    writer.append(PdfReader(io.BytesIO(pdf_bytes())))
    writer.encrypt("private-password")
    encrypted = io.BytesIO()
    writer.write(encrypted)
    for content, code in [
        (encrypted.getvalue(), "ENCRYPTED_DOCUMENT"),
        (pdf_bytes(""), "OCR_REQUIRED"),
        (b"%PDF-1.7 broken", "INVALID_DOCUMENT"),
        (b"not a PDF", "UNSUPPORTED_DOCUMENT"),
    ]:
        with pytest.raises(DomainError, match=code):
            parse_document("file.pdf", content)


def test_docx_zip_bomb_and_xml_entity_rejected():
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"x" * 200000)
    with pytest.raises(DomainError, match="DOCUMENT_TOO_COMPLEX"):
        parse_document("bomb.docx", stream.getvalue())
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr(
            "word/document.xml", '<!DOCTYPE foo [<!ENTITY x SYSTEM "file:///etc/passwd">]>'
        )
    with pytest.raises(DomainError, match="UNSAFE_DOCUMENT_XML"):
        parse_document("entity.docx", stream.getvalue())
    with pytest.raises(DomainError, match="INVALID_DOCUMENT"):
        parse_document("corrupt.docx", b"broken zip")


def test_parser_timeout_and_minimal_child_environment(monkeypatch):
    def timeout(args, **kwargs):
        assert args[1] == "-I"
        assert set(kwargs["env"]) == {"PATH", "LANG"}
        assert kwargs["cwd"] == "/"
        assert kwargs["timeout"] <= 20
        assert kwargs["capture_output"]
        raise subprocess.TimeoutExpired(args, kwargs["timeout"])

    monkeypatch.setattr(subprocess, "run", timeout)
    with pytest.raises(DomainError, match="DOCUMENT_PARSE_TIMEOUT"):
        parse_document("file.pdf", pdf_bytes())


def test_binary_original_version_and_citations(ctx):
    k, owner = ctx["knowledge"], ctx["actors"]["owner"]
    original = docx_bytes()
    first = k.import_document(owner, "rules.docx", original, ["owner", "employee"], "test")
    replay = k.import_document(owner, "rules.docx", original, ["owner", "employee"], "test")
    assert replay["replayed"]
    version = k.get_document(owner, first["document_id"], 1)
    assert version["anchors"][0]["ref"] == "paragraph:1"
    assert k.original_document(owner, first["document_id"], 1)["content"] == original
    evidence = k.search(ctx["actors"]["employee"], "STEEL-01 125.00")
    matching = [item for item in evidence if item["source_id"] == first["document_id"]]
    assert matching and matching[0]["fragment_ref"] == "table:1:row:2"
    updated = docx_bytes("Обновлённый лимит 200 000 рублей.")
    second = k.import_document(owner, "rules.docx", updated, ["owner", "employee"], "test")
    assert second["version"] == 2
    assert k.get_document(owner, first["document_id"], 1)["status"] == "superseded"
    assert k.original_document(owner, first["document_id"], 1)["content"] == original
    k.update_acl(owner, first["document_id"], DocumentACL(roles=["owner"]), "test")
    with pytest.raises(DomainError, match="NOT_FOUND"):
        k.original_document(ctx["actors"]["employee"], first["document_id"], 1)


def test_changed_pdf_layout_preserves_distinct_original_even_with_same_text(ctx):
    k, owner = ctx["knowledge"], ctx["actors"]["owner"]
    original = pdf_bytes()
    first = k.import_document(owner, "price.pdf", original, ["owner"], "test")
    writer = PdfWriter()
    writer.append(PdfReader(io.BytesIO(original)))
    writer.add_metadata({"/Title": "New source revision"})
    changed = io.BytesIO()
    writer.write(changed)
    assert (
        parse_document("price.pdf", original)["text"]
        == parse_document("price.pdf", changed.getvalue())["text"]
    )
    second = k.import_document(owner, "price.pdf", changed.getvalue(), ["owner"], "test")
    assert second["version"] == 2
    assert k.original_document(owner, first["document_id"], 1)["content"] == original


def test_parser_and_index_failures_preserve_published_binary(ctx, monkeypatch):
    k, owner = ctx["knowledge"], ctx["actors"]["owner"]
    first = k.import_document(owner, "price.pdf", pdf_bytes(), ["owner"], "test")
    with pytest.raises(DomainError, match="OCR_REQUIRED"):
        k.import_document(owner, "price.pdf", pdf_bytes(""), ["owner"], "test")
    assert k.get_document(owner, first["document_id"])["current_version"] == 1

    def fail(*args, **kwargs):
        raise RuntimeError("storage failed")

    monkeypatch.setattr(k.store, "add_documents", fail)
    with pytest.raises(DomainError, match="INDEXING_FAILED"):
        k.import_document(owner, "price.pdf", pdf_bytes("Updated offer"), ["owner"], "test")
    assert k.get_document(owner, first["document_id"])["current_version"] == 1
    with ctx["engine"].connect() as conn:
        failed = row(
            conn,
            select(versions).where(
                versions.c.document_id == first["document_id"], versions.c.version == 2
            ),
        )
        assert failed["state"] == "failed"


def test_binary_reindex_keeps_original_version_and_anchors(ctx):
    k, owner = ctx["knowledge"], ctx["actors"]["owner"]
    first = k.import_document(owner, "price.pdf", pdf_bytes(), ["owner"], "test")
    source = k.get_document(owner, first["document_id"], 1)
    original = k.original_document(owner, first["document_id"], 1)
    with transaction(ctx["engine"]) as conn:
        conn.execute(
            versions.update().where(versions.c.id == source["source"]["id"]).values(state="pending")
        )
    indexed = k.import_document(
        owner,
        "price.pdf",
        original["content"],
        ["owner"],
        "reindex",
        first["document_id"],
        content_type=original["media_type"],
    )
    assert indexed["version"] == 1
    assert k.get_document(owner, first["document_id"], 1)["anchors"] == source["anchors"]


def test_original_symlink_escape_and_tampering_rejected(ctx, tmp_path):
    k, owner = ctx["knowledge"], ctx["actors"]["owner"]
    first = k.import_document(owner, "price.pdf", pdf_bytes(), ["owner"], "test")
    path = ctx["settings"].data_dir / "documents" / f"{first['document_id']}-1.original.pdf"
    path.write_bytes(b"tampered")
    with pytest.raises(DomainError, match="SOURCE_HASH_MISMATCH"):
        k.original_document(owner, first["document_id"], 1)
    path.unlink()
    outside = tmp_path / "outside"
    outside.write_bytes(b"private")
    path.symlink_to(outside)
    with pytest.raises(DomainError, match="UNSAFE_STORAGE"):
        k.original_document(owner, first["document_id"], 1)


def test_upgrade_from_v1_preserves_legacy_document(tmp_path):
    from alembic import command
    from alembic.config import Config

    from app.config import Settings
    from app.contracts import Actor
    from app.db import digest, engine_for, migrate, uid
    from app.knowledge import Knowledge
    from app.schema_v1 import documents

    settings = Settings(data_dir=tmp_path / "old-data", _env_file=None)
    settings.data_dir.mkdir()
    config = Config("alembic.ini")
    config.set_main_option("sqlalchemy.url", settings.database_url)
    command.upgrade(config, "0001")
    engine = engine_for(settings)
    document_id, version_id = uid(), uid()
    filename = f"{document_id}-1.txt"
    source = settings.data_dir / "documents"
    source.mkdir()
    (source / filename).write_text("Legacy approved policy", encoding="utf-8")
    with transaction(engine) as conn:
        conn.execute(
            documents.insert().values(
                id=document_id,
                organization_id="legacy",
                name="policy.md",
                roles=["owner"],
                revoked=False,
                current_version=1,
            )
        )
        conn.execute(
            versions.insert().values(
                id=version_id,
                document_id=document_id,
                version=1,
                content_hash=digest("Legacy approved policy"),
                state="indexed",
                file_name=filename,
                observed_at="2026-09-01T00:00:00Z",
            )
        )
    migrate(settings)
    migrate(settings)
    knowledge = Knowledge(engine, settings)
    actor = Actor(id="owner", organization_id="legacy", role="owner", team_id="operations")
    doc = knowledge.get_document(actor, document_id, 1)
    assert doc["content"] == "Legacy approved policy"
    assert doc["anchors"] == []
    original = knowledge.original_document(actor, document_id, 1)
    assert original["original_preserved"] is False
    assert original["filename"] == "policy.md.extracted.txt"
    engine.dispose()


def test_binary_size_cap_and_source_filename_controls():
    with pytest.raises(DomainError, match="UPLOAD_TOO_LARGE"):
        parse_document("too-big.pdf", b"x" * (10 * 1024 * 1024 + 1))
    for name in ("report\n.pdf", "../report.pdf", "report\x00.pdf"):
        with pytest.raises(DomainError, match="UNSUPPORTED_DOCUMENT"):
            parse_document(name, pdf_bytes())
