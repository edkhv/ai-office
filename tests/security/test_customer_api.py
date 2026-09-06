"""Customer journey through authenticated API boundaries and a mocked local model."""

import io
import json
from pathlib import Path

import httpx
import pytest
from docx import Document
from pypdf import PdfReader

from app.config import Settings
from app.errors import DomainError
from app.providers import CrewProvider
from tests.conftest import headers
from tests.security.test_document_formats import docx_bytes, pdf_bytes

CSV = Path("examples/catalogs/synthetic-demo.csv").read_bytes()


def test_customer_document_catalog_quote_approval_task_export(ctx, owner_headers, employee_headers):
    client = ctx["client"]
    source = client.post(
        "/api/v1/documents",
        headers=owner_headers,
        files={
            "file": (
                "request.docx",
                docx_bytes("STEEL-01 × 3"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"roles": "owner,manager"},
    )
    assert source.status_code == 200, source.text
    document = source.json()
    assert (
        client.get(
            f"/api/v1/documents/{document['document_id']}/original?version=1",
            headers=employee_headers,
        ).status_code
        == 404
    )
    original = client.get(
        f"/api/v1/documents/{document['document_id']}/original?version=1", headers=owner_headers
    )
    assert original.status_code == 200 and "attachment" in original.headers["content-disposition"]
    assert Document(io.BytesIO(original.content)).paragraphs[0].text == "STEEL-01 × 3"
    catalog = client.post(
        "/api/v1/catalogs", headers=owner_headers, files={"file": ("demo.csv", CSV, "text/csv")}
    )
    assert catalog.status_code == 200, catalog.text
    version = catalog.json()
    assert (
        client.get(
            f"/api/v1/catalogs/{version['catalog_id']}/versions", headers=owner_headers
        ).json()[0]["id"]
        == version["id"]
    )
    request = {
        "catalog_version_id": version["id"],
        "source_document_id": document["document_id"],
        "source_document_version": 1,
    }
    suggestion = client.post(
        "/api/v1/quote-suggestions",
        headers={**owner_headers, "Idempotency-Key": "customer-suggestion"},
        json=request,
    )
    assert suggestion.status_code == 202, suggestion.text
    ctx["work"].process_one()
    result = client.get(suggestion.json()["status_url"], headers=owner_headers).json()
    assert result["state"] == "completed" and result["result"]["lines"][0]["quantity"] == "3"
    draft = {
        **request,
        "title": "КП для Примера",
        "customer": "Синтетический заказчик",
        "lines": result["result"]["lines"],
        "task": {
            "title": "Получить отзыв по КП",
            "team_id": "procurement",
            "assignee_id": "employee",
            "due_at": "2026-09-08T12:00:00Z",
            "acceptance_criteria": "Зафиксирован следующий шаг",
        },
    }
    response = client.post("/api/v1/quotes", headers=owner_headers, json=draft)
    assert response.status_code == 201, response.text
    quote = response.json()
    assert quote["revision"]["snapshot"]["calculation"]["total"] == "360.00"
    quote_id = quote["id"]
    assert (
        client.get(f"/api/v1/quotes/{quote_id}", headers=headers(ctx, "manager")).status_code == 404
    )
    assert (
        client.get(
            f"/api/v1/quotes/{quote_id}/export?format=pdf", headers=employee_headers
        ).status_code
        == 403
    )
    assert client.get("/api/v1/tasks", headers=owner_headers).json() == []
    pending = client.post(
        f"/api/v1/quotes/{quote_id}/propose",
        headers={**owner_headers, "Idempotency-Key": "customer-proposal"},
        json={"version": 1},
    ).json()
    proposal = pending["proposal"]
    decision = {"decision": "approve", "version": 1, "payload_hash": proposal["payload_hash"]}
    assert (
        client.post(
            f"/api/v1/approvals/{proposal['id']}/decision", headers=owner_headers, json=decision
        ).status_code
        == 200
    )
    assert (
        client.post(
            f"/api/v1/approvals/{proposal['id']}/decision", headers=owner_headers, json=decision
        ).status_code
        == 200
    )
    ctx["work"].process_one()
    tasks = client.get(
        "/api/v1/tasks?mine=true&timezone=Europe/Moscow", headers=employee_headers
    ).json()
    assert len(tasks) == 1 and tasks[0]["assignee_id"] == "employee"
    assert (
        client.get("/api/v1/tasks/" + tasks[0]["id"], headers=employee_headers).status_code == 200
    )
    pdf = client.get(f"/api/v1/quotes/{quote_id}/export?format=pdf", headers=owner_headers)
    text = "\n".join(page.extract_text() for page in PdfReader(io.BytesIO(pdf.content)).pages)
    assert "360.00" in text and "DRAFT" not in text
    word = client.get(f"/api/v1/quotes/{quote_id}/export?format=docx", headers=owner_headers)
    assert "DRAFT" not in " ".join(p.text for p in Document(io.BytesIO(word.content)).paragraphs)


def test_customer_catalog_template_roles_and_large_upload(ctx, owner_headers, employee_headers):
    client = ctx["client"]
    for format in ("csv", "xlsx"):
        assert (
            client.get(
                "/api/v1/catalogs/template?format=" + format, headers=employee_headers
            ).status_code
            == 403
        )
        response = client.get("/api/v1/catalogs/template?format=" + format, headers=owner_headers)
        assert response.status_code == 200
        uploaded = client.post(
            "/api/v1/catalogs",
            headers=owner_headers,
            files={"file": ("template." + format, response.content)},
        )
        assert uploaded.status_code == 200, uploaded.text
    large = "sku,name,unit,price_without_vat,vat_percent,currency\n" + "".join(
        f"X-{i},{'a' * 180},pcs,1.00,20,RUB\n" for i in range(800)
    )
    assert len(large) > 131072
    assert (
        client.post(
            "/api/v1/catalogs", headers=owner_headers, files={"file": ("large.csv", large.encode())}
        ).status_code
        == 200
    )
    assert (
        client.post(
            "/api/v1/documents",
            headers=owner_headers,
            files={"file": ("large.txt", b"x" * 131073)},
            data={"roles": "owner"},
        ).status_code
        == 413
    )
    assert (
        client.post(
            "/api/v1/catalogs", headers=employee_headers, files={"file": ("demo.csv", CSV)}
        ).status_code
        == 403
    )


def test_pdf_original_page_reference_api(ctx, owner_headers):
    client = ctx["client"]
    pdf = pdf_bytes("Orchid customer clause: delivery within three days.")
    imported = client.post(
        "/api/v1/documents",
        headers=owner_headers,
        files={"file": ("terms.pdf", pdf, "application/pdf")},
        data={"roles": "owner"},
    ).json()
    doc = client.get("/api/v1/documents/" + imported["document_id"], headers=owner_headers).json()
    assert doc["anchors"][0]["ref"] == "page:1"
    assert client.get(doc["original_url"], headers=owner_headers).content == pdf


def test_crewai_quote_suggestion_real_agent_mocked_transport(ctx):
    calls = []

    def handle(request):
        calls.append(json.loads(request.content))
        return httpx.Response(
            200,
            json={
                "message": {
                    "content": json.dumps(
                        {
                            "lines": [{"sku": "X", "quantity": "3"}],
                            "unresolved": [],
                            "accompanying_text": "Unsent draft",
                        }
                    )
                }
            },
        )

    settings = Settings(
        data_dir=ctx["settings"].data_dir,
        mode="local_ollama",
        embedding_provider="ollama",
        _env_file=None,
    )
    provider = CrewProvider(settings, httpx.MockTransport(handle))
    result = provider.suggest_quote(
        "X × 3", [{"sku": "X", "name": "Example", "unit": "pcs", "price_without_vat": "999"}]
    )
    assert result["lines"][0]["sku"] == "X"
    assert calls and all("tools" not in call for call in calls)
    assert "price_without_vat" not in json.dumps(calls)


def test_model_unknown_sku_and_oversized_context_rejected(ctx, monkeypatch):
    provider = CrewProvider(
        Settings(
            data_dir=ctx["settings"].data_dir,
            mode="local_ollama",
            embedding_provider="ollama",
            _env_file=None,
        )
    )
    monkeypatch.setattr(
        provider,
        "crew_step",
        lambda *args: json.dumps({"lines": [{"sku": "UNKNOWN", "quantity": "1"}]}),
    )
    with pytest.raises(DomainError, match="INVALID_QUOTE_SUGGESTION"):
        provider.suggest_quote("Need an item", [{"sku": "X", "name": "Item", "unit": "pcs"}])
    with pytest.raises(DomainError, match="CATALOG_TOO_LARGE_FOR_SUGGESTION"):
        provider.suggest_quote(
            "Find an item", [{"sku": str(i), "name": "Item", "unit": "pcs"} for i in range(101)]
        )
